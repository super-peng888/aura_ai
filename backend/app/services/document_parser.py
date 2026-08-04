"""Document parser: PyMuPDF text + optional image extraction, VLM, PaddleOCR.

Simple, direct implementation. No abstract base classes or factories.
Parsing is synchronous; async operations (OSS upload, DB writes) are handled
caller-side after parse() returns.
"""

import io
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional

import fitz

from app.config import get_settings
from app.services.usage_service import usage_service

settings = get_settings()
logger = logging.getLogger(__name__)


def _track_vlm_usage(response, model_name: str) -> None:
    """VLM chat.completions 响应 → 用量埋点。"""
    u = getattr(response, "usage", None)
    usage_service.track(
        "vlm",
        model_name,
        scene="parse",
        prompt_tokens=getattr(u, "prompt_tokens", 0) or 0,
        completion_tokens=getattr(u, "completion_tokens", 0) or 0,
        total_tokens=getattr(u, "total_tokens", 0) or 0,
    )


@dataclass
class ParsedPage:
    page_number: int
    text: str
    images: List[dict] = field(default_factory=list)
    tables: List[str] = field(default_factory=list)
    # 标题条目: {"text": str, "level": int, "pos": int（在 page.text 中的字符位置）}
    headings: List[dict] = field(default_factory=list)
    # 内容格式: plain | markdown（markdown 强制按标题切分，代码围栏原子保留）
    content_format: str = "plain"


@dataclass
class ParseStrategyConfig:
    """解析策略配置，由调用方（根据用户策略）提供。

    VLM 多模态解析的模型配置由调用方（document_parse_service 按策略的
    vlm_model_ref 异步解析）注入 vlm_* 字段；未注入时回落系统级 settings
    （settings.VLM_* + DASHSCOPE_API_KEY）。
    """

    parse_mode: str = "pymupdf"          # pymupdf | paddleocr | vlm（历史别名: pymupdf_rich | ocr）
    chunk_size: int = 800
    chunk_overlap: int = 100
    split_method: str = "sentence"       # sentence | token | structured
    extract_images: bool = False
    dimension: int = 1536
    # vlm 模型配置注入（None = 回落 settings）；vlm_model_ref 为策略行的模型引用，
    # 由 document_parse_service 异步解析后填充其余 vlm_* 字段
    vlm_model_ref: Optional[str] = None
    vlm_model: Optional[str] = None
    vlm_base_url: Optional[str] = None
    vlm_api_key: Optional[str] = None
    vlm_detail: Optional[str] = None


@dataclass
class ParseResult:
    """同步解析结果。图片仅包含原始二进制数据，尚未上传 OSS。"""

    doc_id: str
    pages: List[ParsedPage]
    raw_images: List[dict] = field(default_factory=list)
    mode_used: str = "pymupdf"


# ---------------------------------------------------------------------------
# 公共工具
# ---------------------------------------------------------------------------

IMG_PLACEHOLDER_PATTERN = r"\[IMG:([a-zA-Z0-9_]+)\]"


def _simple_split(text: str, chunk_size: int, chunk_overlap: int) -> List[str]:
    """兜底分片：按段落粗分，超过 chunk_size 时截断。"""
    paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
    chunks = []
    current = []
    current_len = 0
    for para in paragraphs:
        if current_len + len(para) > chunk_size and current:
            chunks.append("\n".join(current))
            overlap = current[-1] if len(current[-1]) < chunk_overlap else current[-1][-chunk_overlap:]
            current = [overlap, para]
            current_len = len(overlap) + len(para)
        else:
            current.append(para)
            current_len += len(para) + 1
    if current:
        chunks.append("\n".join(current))
    return chunks or [text]


def _sentence_split(text: str, chunk_size: int, chunk_overlap: int) -> List[str]:
    """使用 llama_index SentenceSplitter 做句子级分片。"""
    try:
        from llama_index.core.node_parser import SentenceSplitter

        splitter = SentenceSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        return splitter.split_text(text)
    except ImportError:
        return _simple_split(text, chunk_size, chunk_overlap)


def _token_split(text: str, chunk_size: int, chunk_overlap: int) -> List[str]:
    """按 token 估算分片（1 token ≈ 0.75 中文字符 或 4 英文字符）。"""
    # 简化实现：按字符长度分，假设平均每个 token 2 个字符
    token_size = max(chunk_size * 2, 1)
    # overlap >= size 时窗口不再前进会死循环，强制 clamp
    overlap_size = min(chunk_overlap * 2, token_size - 1)
    if len(text) <= token_size:
        return [text]
    chunks = []
    start = 0
    while start < len(text):
        end = start + token_size
        chunks.append(text[start:end])
        start = end - overlap_size
        if start >= len(text):
            break
    return chunks or [text]


def _structured_split(text: str, _chunk_size: int, _chunk_overlap: int) -> List[str]:
    """结构化分片：按标题/章节分割（简单实现：按 # 或 数字标题）。"""
    # 匹配 Markdown 标题或中文/数字章节标题
    pattern = r"(?:\n|^)(#{1,6}\s+|第[一二三四五六七八九十\d]+章[\s、]|\d+\.\s+\S|【.+?】)(?=\n)"
    parts = re.split(pattern, text)
    if len(parts) <= 2:
        return _simple_split(text, _chunk_size, _chunk_overlap)
    chunks = []
    current = parts[0] if parts[0] else ""
    i = 1
    while i < len(parts):
        header = parts[i]
        body = parts[i + 1] if i + 1 < len(parts) else ""
        candidate = current + "\n" + header + body
        if len(candidate) > _chunk_size and current:
            chunks.append(current.strip())
            current = header + body
        else:
            current = candidate
        i += 2
    if current.strip():
        chunks.append(current.strip())
    return chunks or [text]


SPLIT_HANDLERS = {
    "sentence": _sentence_split,
    "token": _token_split,
    "structured": _structured_split,
}


# ---------------------------------------------------------------------------
# Markdown 按标题切分（代码围栏原子保留）
# ---------------------------------------------------------------------------

_MD_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")


def _extract_markdown_headings(text: str) -> List[dict]:
    """提取 Markdown 标题（代码围栏内的 # 行不算标题）。"""
    headings = []
    fence = None
    pos = 0
    for line in text.splitlines(keepends=True):
        stripped = line.lstrip()
        if fence:
            if stripped.startswith(fence):
                fence = None
        elif stripped.startswith("```") or stripped.startswith("~~~"):
            fence = stripped[:3]
        else:
            m = _MD_HEADING_RE.match(line.rstrip("\n"))
            if m:
                headings.append({"text": m.group(2).strip(), "level": len(m.group(1)), "pos": pos})
        pos += len(line)
    return headings


def _markdown_sections(text: str) -> List[tuple]:
    """按标题行切分为节（代码围栏内的 # 不作为边界）。Returns [(pos, section_text)]。"""
    sections = []
    fence = None
    cur_start = 0
    pos = 0
    for line in text.splitlines(keepends=True):
        stripped = line.lstrip()
        if fence:
            if stripped.startswith(fence):
                fence = None
        elif stripped.startswith("```") or stripped.startswith("~~~"):
            fence = stripped[:3]
        elif _MD_HEADING_RE.match(line.rstrip("\n")) and pos > cur_start:
            section = text[cur_start:pos]
            if section.strip():
                sections.append((cur_start, section))
            cur_start = pos
        pos += len(line)
    tail = text[cur_start:]
    if tail.strip():
        sections.append((cur_start, tail))
    return sections or [(0, text)]


def _split_section_blocks(section: str) -> List[str]:
    """把节拆为原子块：代码围栏整体为一块，其余按空行分段。"""
    blocks = []
    cur: List[str] = []
    fence = None
    for line in section.splitlines(keepends=True):
        stripped = line.lstrip()
        if fence:
            cur.append(line)
            if stripped.startswith(fence):
                blocks.append("".join(cur))
                cur = []
                fence = None
            continue
        if stripped.startswith("```") or stripped.startswith("~~~"):
            if cur:
                blocks.append("".join(cur))
                cur = []
            fence = stripped[:3]
            cur.append(line)
            continue
        if not line.strip():
            if cur:
                blocks.append("".join(cur))
                cur = []
            continue
        cur.append(line)
    if cur:
        blocks.append("".join(cur))
    return blocks


def _is_table_block(blk: str) -> bool:
    """判断块是否为 Markdown 表格（GFM）：含一行分隔行（如 | --- | --- |）。

    表格行间无空行，_split_section_blocks 已整体保为一块；此处用于让
    超长表格与代码围栏同等原子，不被 _simple_split 按行切散（表头与数据行分离）。
    """
    for line in blk.splitlines():
        s = line.strip()
        if "|" in s and "-" in s and set(s) <= set("|:- "):
            return True
    return False


# 父块组上限（字符）：限制兄弟节聚合后的父块体积，兼顾 LLM 上下文预算
# 与 Milvus metadata JSON 字段容量（64KB，6000 汉字约 18KB UTF-8）。
_PARENT_GROUP_MAX_CHARS = 6000


def _section_heading_level(section: str) -> Optional[int]:
    """返回节首行标题级别（1~6）；无标题（导言节）返回 None。"""
    first_line = section.lstrip("\n").split("\n", 1)[0]
    m = _MD_HEADING_RE.match(first_line.rstrip())
    return len(m.group(1)) if m else None


def _aggregate_section_children(section: str, start: int, child_size: int) -> List[tuple]:
    """将单个标题节聚合为子块 (pos, child_text) 序列。

    - 节 <= child_size 时整节为一个子块；
    - 超长节按空行块聚合，代码围栏与表格永远作为原子块不被切开。
    """
    if len(section) <= child_size:
        return [(start, section.rstrip())]
    pieces = []
    offset = start
    cur_parts: List[str] = []
    cur_len = 0
    cur_start = start
    for block in _split_section_blocks(section):
        blk = block.rstrip("\n")
        if cur_parts and cur_len + len(blk) > child_size:
            pieces.append((cur_start, "\n\n".join(cur_parts)))
            cur_parts, cur_len = [], 0
            cur_start = offset
        atomic = blk.lstrip().startswith(("```", "~~~")) or _is_table_block(blk)
        if not cur_parts and len(blk) > child_size and not atomic:
            # 普通超长段落退化细分；代码围栏/表格保持原子
            sub_off = offset
            for sub in _simple_split(blk, child_size, 0):
                pieces.append((sub_off, sub))
                sub_off += len(sub)
            cur_start = offset + len(block)
        else:
            cur_parts.append(blk)
            cur_len += len(blk) + 2
        offset += len(block)
    if cur_parts:
        pieces.append((cur_start, "\n\n".join(cur_parts)))
    return pieces or [(start, section.rstrip())]


def _markdown_parents(text: str, child_size: int) -> List[dict]:
    """Markdown 按标题切分为父子结构（small-to-big 检索）。

    父块＝连续的「同级或更深级」标题节贪心聚合而成的组（遇更高级标题或
    超过 _PARENT_GROUP_MAX_CHARS 才断组）：典型如「报文代码块节 + 后续 N 个
    同级数据元详解节」合为一个父块，命中任意子块可整组带回，避免兄弟节
    内容残缺。组内各节再细分为子块（精确检索，代码围栏/表格原子），
    命中子块后回捞整组内容送 LLM。

    无标题的导言（preamble）独立成组，不吸收后续节。

    Returns:
        list of {"parent_idx", "parent_pos", "parent_text", "children": [(pos, text)]}
    """
    groups: List[List[tuple]] = []
    anchor_level: Optional[int] = None
    group_len = 0
    for start, section in _markdown_sections(text):
        level = _section_heading_level(section)
        new_group = (
            not groups
            or level is None  # 导言只会是首节，防御性判断
            or anchor_level is None  # 上一组锚在导言，不吸收
            or level < anchor_level  # 更高级标题断组
            or group_len + len(section) > _PARENT_GROUP_MAX_CHARS
        )
        if new_group:
            groups.append([(start, section)])
            anchor_level = level
            group_len = len(section)
        else:
            groups[-1].append((start, section))
            group_len += len(section)

    parents = []
    for pidx, group in enumerate(groups):
        children: List[tuple] = []
        for s_start, s_text in group:
            children.extend(_aggregate_section_children(s_text, s_start, child_size))
        parents.append({
            "parent_idx": pidx,
            "parent_pos": group[0][0],
            "parent_text": "".join(s for _, s in group).rstrip(),
            "children": children,
        })
    return parents or [{"parent_idx": 0, "parent_pos": 0, "parent_text": text.rstrip(),
                        "children": [(0, text.rstrip())]}]


def _markdown_split(text: str, chunk_size: int) -> List[tuple]:
    """Markdown 按标题切分为扁平 (pos, chunk_text) 序列（不含父块信息）。

    - 标题（# ~ ######，代码围栏内除外）是切分边界，节内容完整保留；
    - 超长节按空行分块聚合，代码围栏永远作为原子块不被切开；
    - 标题边界即语义边界，不做 overlap。
    """
    pieces = []
    for parent in _markdown_parents(text, chunk_size):
        pieces.extend(parent["children"])
    return pieces or [(0, text)]


def _split_by_page_headings(text: str, page_headings: List[dict], strategy: ParseStrategyConfig) -> List[tuple]:
    """按提取到的真实标题位置切分页面文本。

    Returns:
        list of (start_pos, chunk_text)；超过 chunk_size 的节用 _simple_split 再细分。
    """
    positions = sorted({h.get("pos", 0) for h in page_headings if 0 < h.get("pos", 0) < len(text)})
    boundaries = [0] + positions + [len(text)]
    pieces = []
    for start, end in zip(boundaries, boundaries[1:]):
        section = text[start:end].strip()
        if not section:
            continue
        if len(section) <= strategy.chunk_size:
            pieces.append((start, section))
        else:
            offset = start
            for sub in _simple_split(section, strategy.chunk_size, strategy.chunk_overlap):
                pieces.append((offset, sub))
                offset += len(sub)
    return pieces


def split_pages_to_chunks(pages: List[ParsedPage], strategy: ParseStrategyConfig, doc_id: str) -> List[dict]:
    """将解析后的页面分片。

    chunk dict 既有字段（chunk_id/doc_id/page/content/image_ids）保持不变，
    新增可选字段：
    - heading: 当前 chunk 所属的标题路径（如 "第一章 > 1.2 节"），无标题为空串
    - chunk_type: "text"（默认）或 "table"（Markdown 表格独立成块）
    """
    splitter = SPLIT_HANDLERS.get(strategy.split_method, _sentence_split)
    chunks = []
    chunk_index = 0
    heading_stack: List[dict] = []  # 跨页维护的标题层级栈 [{"level": int, "text": str}]

    def heading_path() -> str:
        return " > ".join(h["text"] for h in heading_stack)

    for page in pages:
        text = page.text or ""
        piece_parents: List[Optional[tuple]] = []  # 与 pieces 对齐：每个子块的 (parent_id, parent_text) 或 None
        page_headings = sorted(page.headings or [], key=lambda h: h.get("pos", 0))
        # 页面级图片（VLM 整页分析等无占位符模式）：整页文本中没有任何 IMG
        # 占位符时，将 page.images 作为页级引用附加到该页每个 chunk，
        # 供检索侧按页引用整页截图；有占位符时按占位符精确定位，不重复附加。
        page_image_ids = [img["image_id"] for img in (page.images or []) if img.get("image_id")]
        page_has_placeholders = bool(re.search(IMG_PLACEHOLDER_PATTERN, text))

        if getattr(page, "content_format", "plain") == "markdown":
            # Markdown 强制按标题切分（代码围栏原子保留），无视 split_method：
            # sentence/token 等通用切分会把完整代码块/报文切碎，检索命中的
            # 片段残缺会误导 LLM 基于半截内容自由发挥。
            # 父子分块：子块精确检索，多子块的父节内容回捞送 LLM（见
            # rag_pipeline._expand_to_parents）；单子块节无需父块（子即父）。
            pieces = []
            for parent in _markdown_parents(text, strategy.chunk_size):
                children = parent["children"]
                parent_meta = None
                if len(children) > 1:
                    parent_meta = (f"{doc_id}_parent_{parent['parent_idx']:04d}", parent["parent_text"])
                for cpos, ctext in children:
                    pieces.append((cpos, ctext))
                    piece_parents.append(parent_meta)
        elif strategy.split_method == "structured" and page_headings:
            # structured 模式优先使用真实标题切分
            pieces = _split_by_page_headings(text, page_headings, strategy)
        else:
            pieces = []
            cursor = 0
            for chunk_text in splitter(text, strategy.chunk_size, strategy.chunk_overlap):
                # 尽力定位 chunk 在页面文本中的起始位置，用于挂标题
                pos = text.find(chunk_text[:30], cursor) if chunk_text else -1
                if pos == -1:
                    pos = cursor
                pieces.append((pos, chunk_text))
                cursor = pos + max(len(chunk_text), 1)

        applied_idx = 0
        for piece_idx, (pos, chunk_text) in enumerate(pieces):
            # 应用位于该 chunk 之前的标题，维护层级栈
            while applied_idx < len(page_headings) and page_headings[applied_idx].get("pos", 0) <= pos:
                h = page_headings[applied_idx]
                while heading_stack and heading_stack[-1]["level"] >= h["level"]:
                    heading_stack.pop()
                heading_stack.append(h)
                applied_idx += 1
            img_ids = re.findall(IMG_PLACEHOLDER_PATTERN, chunk_text)
            if page_image_ids and not page_has_placeholders:
                img_ids = list(dict.fromkeys([*img_ids, *page_image_ids]))
            chunk = {
                "chunk_id": f"{doc_id}_chunk_{chunk_index:04d}",
                "doc_id": doc_id,
                "page": page.page_number,
                "content": chunk_text,
                "image_ids": img_ids,
                "heading": heading_path(),
                "chunk_type": "text",
            }
            parent_meta = piece_parents[piece_idx] if piece_idx < len(piece_parents) else None
            if parent_meta:
                chunk["parent_id"], chunk["parent_content"] = parent_meta
            chunks.append(chunk)
            chunk_index += 1

        # 应用页面剩余标题，作为表格 chunk 的上下文
        while applied_idx < len(page_headings):
            h = page_headings[applied_idx]
            while heading_stack and heading_stack[-1]["level"] >= h["level"]:
                heading_stack.pop()
            heading_stack.append(h)
            applied_idx += 1

        for table_md in page.tables or []:
            path = heading_path()
            content = f"{path}\n{table_md}" if path else table_md
            chunks.append({
                "chunk_id": f"{doc_id}_chunk_{chunk_index:04d}",
                "doc_id": doc_id,
                "page": page.page_number,
                "content": content,
                "image_ids": [],
                "heading": path,
                "chunk_type": "table",
            })
            chunk_index += 1
    return chunks


# ---------------------------------------------------------------------------
# 图片提取（纯同步，仅返回二进制数据，不调用异步 OSS 上传）
# ---------------------------------------------------------------------------

_IMAGE_RENDER_ZOOM = 2.0        # 渲染缩放下限（低分辨率原图的清晰度兜底）
_IMAGE_MAX_RENDER_ZOOM = 6.0    # 渲染缩放上限（≈432dpi，防止超高分辨率原图产出过大文件）
_IMAGE_MIN_SIZE_PT = 50.0   # 显示尺寸过滤（pt）：过小的装饰图标/线条跳过
_IMAGE_MERGE_GAP_PT = 5.0   # 矩形合并间距（pt）：被切成条带的大图垂直相邻，合并复原


def _merge_image_rects(rects: List["fitz.Rect"], gap: float) -> List["fitz.Rect"]:
    """合并重叠或近邻（间距 ≤ gap）的矩形，迭代至收敛。

    Word/扫描件导出的 PDF 常把一张视觉大图切成多个条带 xref，
    逐个提取只能拿到残片；按显示矩形合并后再渲染可复原完整图。
    """
    merged = [fitz.Rect(r) for r in rects]
    changed = True
    while changed:
        changed = False
        out: List[fitz.Rect] = []
        for r in merged:
            hit = None
            for m in out:
                expanded = fitz.Rect(m.x0 - gap, m.y0 - gap, m.x1 + gap, m.y1 + gap)
                if expanded.intersects(r):
                    hit = m
                    break
            if hit is not None:
                hit.include_rect(r)
                changed = True
            else:
                out.append(fitz.Rect(r))
        merged = out
    return merged


def _extract_images_raw(doc: fitz.Document, doc_id: str) -> List[dict]:
    """按显示区域渲染提取图片：get_image_rects 定位 + get_pixmap(clip) 渲染。

    旧版用 extract_image(xref) 直取内嵌位图，丢失 PDF 变换矩阵（镜像/旋转
    错乱）且分条带存储的大图只能取到残片（截断）；改为在页面上按显示
    矩形渲染，并合并重叠/近邻矩形复原条带图，所见即所得。

    每张图片附带 y_pos（在页面中的纵向坐标，供占位符按阅读顺序插入）。
    渲染缩放按内嵌图原生像素密度自适应（原生像素 ÷ 显示 pt），尽量还原
    原始分辨率，避免固定倍数导致高清扫描件/截图发虚。
    """
    raw_images = []
    for page_idx in range(len(doc)):
        page = doc[page_idx]
        page_number = page_idx + 1
        rects: List[fitz.Rect] = []
        rect_zooms: List[float] = []
        for img_info in page.get_images(full=True):
            xref, native_w, native_h = img_info[0], img_info[2], img_info[3]
            try:
                for r in page.get_image_rects(xref):
                    rr = fitz.Rect(r)
                    rr.intersect(page.rect)  # 裁剪到页面内，避免超出边界的空白
                    if not rr.is_empty:
                        rects.append(rr)
                        # 还原该内嵌图原生分辨率所需的渲染缩放
                        rect_zooms.append(max(native_w / rr.width, native_h / rr.height))
            except Exception:
                continue

        regions = [
            r for r in _merge_image_rects(rects, _IMAGE_MERGE_GAP_PT)
            if r.width >= _IMAGE_MIN_SIZE_PT and r.height >= _IMAGE_MIN_SIZE_PT
        ]
        regions.sort(key=lambda r: (r.y0, r.x0))

        for seq, rect in enumerate(regions, start=1):
            # 区域内所有内嵌图的最高原生密度决定缩放，2x 保底、上限封顶
            zoom = _IMAGE_RENDER_ZOOM
            for rr, z in zip(rects, rect_zooms):
                if rect.intersects(rr):
                    zoom = max(zoom, z)
            zoom = min(zoom, _IMAGE_MAX_RENDER_ZOOM)
            try:
                pix = page.get_pixmap(clip=rect, matrix=fitz.Matrix(zoom, zoom))
                data = pix.tobytes("png")
            except Exception:
                continue

            image_id = f"{doc_id}_p{page_number}_{seq:02d}"
            raw_images.append({
                "image_id": image_id,
                "data": data,
                "content_type": "image/png",
                "width": pix.width,
                "height": pix.height,
                "page_number": page_number,
                "seq": seq,
                "y_pos": rect.y0,
                "filename": f"{image_id}.png",
            })
    return raw_images


def _assemble_page_text(page: fitz.Page, text_blocks: list, images: List[dict], page_number: int) -> str:
    """组装页面文本：按 y 坐标排序，在图片真实位置插入占位符。

    图片取提取时记录的 y_pos；定位失败时追加到页尾（保持阅读顺序），
    不再像旧版那样全部堆到页首。
    """
    contents = []
    for block in text_blocks:
        if block[6] == 0:  # text block
            y_pos = block[1]
            text = block[4].strip()
            if text:
                contents.append({"type": "text", "y_pos": y_pos, "content": text})

    page_bottom = page.rect.height + 1 if page is not None else float("inf")
    for img in images:
        y_pos = img.get("y_pos")
        contents.append({
            "type": "image",
            "y_pos": y_pos if y_pos is not None else page_bottom,
            "content": f"\n[IMG:{img['image_id']}]\n",
        })

    contents.sort(key=lambda x: x["y_pos"])
    return "\n\n".join(c["content"] for c in contents)


def _extract_tables(page: fitz.Page) -> List[str]:
    """提取页面中的表格（Markdown 格式）。"""
    tables = []
    try:
        tab = page.find_tables()
        for t in tab.tables:
            tables.append(t.to_pandas().to_markdown())
    except Exception:
        pass
    return tables


_BOLD_FLAG = 16  # PyMuPDF span flags: 2**4 = bold
_HEADING_MAX_LEN = 60
_HEADING_MIN_RATIO = 1.15   # 字号 >= 正文 * 1.15 视为标题
_HEADING_LEVEL2_RATIO = 1.25
_HEADING_LEVEL1_RATIO = 1.40
# 以句读结尾的短行视为正文而非标题
_HEADING_EXCLUDE_ENDINGS = "。；，,.;!！?？"


def _extract_headings(page: fitz.Page) -> List[dict]:
    """从页面提取标题（启发式，纯 PyMuPDF span 样式）。

    规则：
    - 按字符数加权统计字号，出现最多的字号视为正文字号；
    - 短行（<=60 字符、不以句读结尾）满足以下任一条件视为标题：
      - 行内最大字号 >= 正文字号 * 1.15；
      - 行内全部 span 加粗且字号 >= 正文字号；
    - 层级按字号比例划分：>=1.40 -> 1，>=1.25 -> 2，否则 -> 3。

    无法判断（扫描件/无样式）时返回空列表，行为与旧版一致。
    """
    try:
        data = page.get_text("dict")
    except Exception:
        return []

    size_weights: dict = {}
    lines = []
    for block in data.get("blocks", []):
        if block.get("type") != 0:
            continue
        for line in block.get("lines", []):
            spans = line.get("spans", [])
            text = "".join(s.get("text", "") for s in spans).strip()
            if not text:
                continue
            max_size = max((s.get("size", 0) for s in spans), default=0)
            all_bold = bool(spans) and all(
                (s.get("flags", 0) & _BOLD_FLAG) or "bold" in s.get("font", "").lower()
                for s in spans
            )
            lines.append({"text": text, "size": max_size, "bold": all_bold})
            for s in spans:
                sz = round(s.get("size", 0), 1)
                size_weights[sz] = size_weights.get(sz, 0) + len(s.get("text", ""))

    if not lines or not size_weights:
        return []

    body_size = max(size_weights, key=size_weights.get)
    if not body_size:
        return []

    headings = []
    for line in lines:
        text = line["text"]
        if len(text) > _HEADING_MAX_LEN or text.endswith(tuple(_HEADING_EXCLUDE_ENDINGS)):
            continue
        ratio = line["size"] / body_size
        if ratio >= _HEADING_MIN_RATIO or (line["bold"] and line["size"] >= body_size):
            if ratio >= _HEADING_LEVEL1_RATIO:
                level = 1
            elif ratio >= _HEADING_LEVEL2_RATIO:
                level = 2
            else:
                level = 3
            headings.append({"text": text, "level": level})
    return headings


def _locate_headings(headings: List[dict], page_text: str) -> List[dict]:
    """为标题条目标注其在 page.text 中的字符位置（pos 字段，尽力而为）。"""
    located = []
    search_from = 0
    for h in headings:
        pos = page_text.find(h["text"], search_from)
        if pos == -1:
            pos = page_text.find(h["text"])
        if pos == -1:
            continue
        located.append({**h, "pos": pos})
        search_from = pos + len(h["text"])
    return located


# ---------------------------------------------------------------------------
# 主解析器
# ---------------------------------------------------------------------------

# 纯文本资源扩展名：按扩展名自动路由到纯文本解析，无需选择解析模式
_TEXT_FILE_EXTS = {".txt", ".md", ".markdown", ".json", ".csv", ".log"}

_IMAGE_FILE_EXTS = (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp")

# Word 文档扩展名：PyMuPDF 不支持 doc/docx，自动路由到 python-docx 转写。
# .doc 也先尝试按 docx 打开：Windows 上 mimetypes 可能把 docx 的 mime 猜成
# .doc 后缀，真 docx 内容仍可正常解析；真老格式才报错指引。
_DOCX_FILE_EXTS = (".docx", ".doc")


def _read_text_with_fallback(file_path: str) -> str:
    """读取纯文本文件：优先 UTF-8，失败回退 GBK，最后容错替换。"""
    with open(file_path, "rb") as f:
        data = f.read()
    for encoding in ("utf-8", "gbk"):
        try:
            return data.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return data.decode("utf-8", errors="replace")


def _parse_text_file(file_path: str, doc_id: str, strategy: ParseStrategyConfig) -> ParseResult:
    """纯文本解析（txt/md/json/csv 等）：整篇读入为单页，交由分块器切分。

    Markdown 文件额外提取标题并标记 content_format=markdown，
    分块时强制按标题切分（见 split_pages_to_chunks）。
    """
    text = _read_text_with_fallback(file_path)
    is_md = Path(file_path).suffix.lower() in (".md", ".markdown")
    page = ParsedPage(
        page_number=1,
        text=text,
        headings=_extract_markdown_headings(text) if is_md else [],
        content_format="markdown" if is_md else "plain",
    )
    return ParseResult(
        doc_id=doc_id,
        pages=[page],
        raw_images=[],
        mode_used="text",
    )


# ---------------------------------------------------------------------------
# DOCX 解析（python-docx 转写为 Markdown，接入标题切分管线）
# ---------------------------------------------------------------------------

_DOCX_IMAGE_SUFFIXES = {
    "image/png": ".png", "image/jpeg": ".jpg", "image/gif": ".gif",
    "image/bmp": ".bmp", "image/webp": ".webp",
}

_DOCX_HEADING_RE = re.compile(r"heading\s*(\d+)", re.IGNORECASE)


def _docx_heading_level(para) -> int:
    """Word 段落样式 -> Markdown 标题级别（1-6）；非标题返回 0。"""
    try:
        name = para.style.name or ""
    except Exception:
        return 0
    m = _DOCX_HEADING_RE.match(name.strip())
    if m:
        return min(int(m.group(1)), 6)
    return 1 if name.strip().lower() == "title" else 0


def _docx_table_to_markdown(table) -> str:
    """Word 表格 -> Markdown 表格（合并单元格文本重复展开，可接受）。"""
    rows = []
    for r in table.rows:
        cells = [" ".join(c.text.split()).replace("|", "\\|") for c in r.cells]
        rows.append("| " + " | ".join(cells) + " |")
    if not rows:
        return ""
    sep = "| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |"
    return "\n".join([rows[0], sep, *rows[1:]])


def _parse_docx_file(file_path: str, doc_id: str, strategy: ParseStrategyConfig) -> ParseResult:
    """DOCX 解析：按 body 顺序转写为 Markdown，接入标题切分管线。

    - 标题样式（Heading N / Title）-> #/##/### 标题，保留层级语义；
    - 表格 -> Markdown 表格（切分时按标题路径挂接）；
    - 内嵌图片 -> 提取二进制 + 段落原位 [IMG:] 占位符（页眉页脚天然省略）；
    - Word 无固定分页概念，整篇作为单页，分块器按标题切分不受影响。

    所有 parse_mode 均走此路径：PyMuPDF/OCR/VLM 都无法直接打开 Word 格式。
    """
    from docx import Document as DocxDocument
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    try:
        docx = DocxDocument(file_path)
    except Exception as e:
        raise ValueError(
            "无法解析该 Word 文档：.doc 老格式请先另存为 .docx 再上传"
        ) from e

    lines: List[str] = []
    raw_images: List[dict] = []
    seq = 0
    for child in docx.element.body.iterchildren():
        if child.tag == qn("w:p"):
            para = Paragraph(child, docx)
            text = para.text.strip()
            level = _docx_heading_level(para)
            if text:
                lines.append(f"{'#' * level} {text}" if level else text)
            # 段内图片：提取二进制并在段落之后插入占位符
            for blip in child.iter(qn("a:blip")):
                rid = blip.get(qn("r:embed"))
                part = docx.part.related_parts.get(rid) if rid else None
                data = getattr(part, "blob", None)
                if not data:
                    continue
                seq += 1
                image_id = f"{doc_id}_p1_{seq:02d}"
                content_type = getattr(part, "content_type", "") or "image/png"
                suffix = _DOCX_IMAGE_SUFFIXES.get(content_type, ".png")
                raw_images.append({
                    "image_id": image_id,
                    "data": data,
                    "content_type": content_type,
                    "width": None,
                    "height": None,
                    "page_number": 1,
                    "seq": seq,
                    "filename": f"{image_id}{suffix}",
                })
                lines.append(f"[IMG:{image_id}]")
        elif child.tag == qn("w:tbl"):
            md = _docx_table_to_markdown(Table(child, docx))
            if md:
                lines.append(md)

    text = "\n\n".join(lines)
    page = ParsedPage(
        page_number=1,
        text=text,
        images=[{"image_id": img["image_id"]} for img in raw_images],
        headings=_extract_markdown_headings(text),
        content_format="markdown",
    )
    return ParseResult(
        doc_id=doc_id,
        pages=[page],
        raw_images=raw_images,
        mode_used="docx",
    )


def parse_document(file_path: str, doc_id: str, strategy: ParseStrategyConfig) -> ParseResult:
    """同步解析文档。不调用任何异步 I/O。

    路由规则：
    - 纯文本扩展名（txt/md/json/csv/log）：自动走纯文本解析，无视 parse_mode；
    - 图片扩展名：paddleocr/ocr 模式直接 OCR，其余模式走 VLM 描述；
    - Word 扩展名（docx/doc）：python-docx 转写为 Markdown，无视 parse_mode；
    - 其余（PDF 等）：按 parse_mode 选择 pymupdf / paddleocr / vlm，
      历史别名 pymupdf_rich 归一为 pymupdf+extract_images，ocr 归一为 paddleocr。

    Returns:
        ParseResult: 包含解析后的页面、原始图片数据（未上传）、解析模式。
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext in _TEXT_FILE_EXTS:
        return _parse_text_file(file_path, doc_id, strategy)

    if ext in _IMAGE_FILE_EXTS:
        return _parse_image(file_path, doc_id, strategy)

    if ext in _DOCX_FILE_EXTS:
        return _parse_docx_file(file_path, doc_id, strategy)

    # 历史别名归一（直接改写 strategy，保证 mode_used 与落库 parse_mode 为规范值）
    if strategy.parse_mode == "pymupdf_rich":
        strategy.parse_mode = "pymupdf"
        strategy.extract_images = True

    if strategy.parse_mode == "vlm":
        return _parse_with_vlm(file_path, doc_id, strategy)

    if strategy.parse_mode in ("paddleocr", "ocr"):
        return _parse_with_paddleocr(file_path, doc_id, strategy)

    return _parse_with_pymupdf(file_path, doc_id, strategy)


def _parse_with_pymupdf(file_path: str, doc_id: str, strategy: ParseStrategyConfig) -> ParseResult:
    """PyMuPDF 解析：文本 + 可选图片提取。"""
    doc = fitz.open(file_path)
    pages = []
    raw_images = []
    try:
        # 一次性提取全文档图片（旧版在页循环内反复全量提取，复杂度 O(页数×图片数)）
        all_images = _extract_images_raw(doc, doc_id) if strategy.extract_images else []
        for page_idx in range(len(doc)):
            page = doc[page_idx]
            page_number = page_idx + 1

            text_blocks = page.get_text("blocks", sort=True)
            # 只保留当前页的图片
            images = [img for img in all_images if img["page_number"] == page_number]
            raw_images.extend(images)

            page_text = _assemble_page_text(page, text_blocks, images, page_number)
            tables = _extract_tables(page)
            headings = _locate_headings(_extract_headings(page), page_text)

            pages.append(ParsedPage(
                page_number=page_number,
                text=page_text,
                images=[{"image_id": img["image_id"]} for img in images],
                tables=tables,
                headings=headings,
            ))
    finally:
        doc.close()

    return ParseResult(
        doc_id=doc_id,
        pages=pages,
        raw_images=raw_images,
        mode_used=strategy.parse_mode,
    )


def _resolve_vlm_config(strategy: Optional[ParseStrategyConfig] = None) -> dict:
    """解析生效 VLM 配置：优先用策略注入值，缺省回落系统级 settings。"""
    s = strategy
    return {
        "model": (s.vlm_model if s else None) or settings.VLM_MODEL,
        "base_url": (s.vlm_base_url if s else None) or settings.VLM_API_BASE,
        "api_key": (s.vlm_api_key if s else None) or settings.DASHSCOPE_API_KEY,
        "detail": (s.vlm_detail if s else None) or getattr(settings, "VLM_DETAIL_LEVEL", "high") or "high",
        "max_tokens": _VLM_MAX_TOKENS,
    }


def _create_vlm_client(strategy: Optional[ParseStrategyConfig] = None):
    """创建 VLM OpenAI 兼容客户端；未配置 API Key 时抛出带配置指引的错误。"""
    from openai import OpenAI

    vlm = _resolve_vlm_config(strategy)
    if not vlm["api_key"]:
        raise ValueError(
            "VLM 多模态解析未配置 API Key，请在模型配置页配置或设置系统环境变量 DASHSCOPE_API_KEY"
        )
    client = OpenAI(api_key=vlm["api_key"], base_url=vlm["base_url"])
    return client, vlm


_VLM_MAX_TOKENS = 4096

_VLM_PAGE_PROMPT = """你是专业的文档解析引擎。请把这页文档图片完整转写为规范的 Markdown：
1. 标题：按文档真实层级输出 #/##/### 标题，普通正文不要误标为标题；
2. 表格：用 Markdown 表格语法完整输出所有单元格内容；
3. 图片/图表/流程图：在其原位置用一段文字客观描述内容（含关键数据与流向）；
4. 正文：保持原文文字，不改写、不总结、不遗漏；页眉、页脚、页码省略；
5. 只输出 Markdown 内容本身，不要用代码围栏包裹整页输出，不要添加任何解释。"""


def _strip_md_fence(text: str) -> str:
    """剥掉模型偶尔包裹整页输出的 ```markdown 围栏（仅整体包裹时）。"""
    m = re.match(r"^```[a-zA-Z]*\n([\s\S]*?)\n?```$", text.strip())
    return m.group(1) if m else text


def _vlm_transcribe_page(client, vlm: dict, page: "fitz.Page", heading_stack: List[dict]) -> str:
    """单页截图 → VLM 转写为结构化 Markdown（base64 data URI 发送）。

    heading_stack 为已解析页的标题层级栈，注入 prompt 保证跨页标题层级连贯。"""
    import base64

    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
    page_b64 = base64.b64encode(pix.tobytes("png")).decode()
    prompt = _VLM_PAGE_PROMPT
    if heading_stack:
        path = " > ".join(h["text"] for h in heading_stack)
        prompt += f"\n\n上下文：本页之前的文档标题层级为「{path}」，请据此延续判断本页标题的级别。"
    response = client.chat.completions.create(
        model=vlm["model"],
        messages=[{
            "role": "user",
            "content": [
                {"type": "text", "text": prompt},
                {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{page_b64}", "detail": vlm["detail"]}},
            ],
        }],
        max_tokens=vlm["max_tokens"],
    )
    _track_vlm_usage(response, vlm["model"])
    return _strip_md_fence((response.choices[0].message.content or "").strip())


def _parse_with_vlm(file_path: str, doc_id: str, strategy: ParseStrategyConfig) -> ParseResult:
    """VLM 多模态解析：每页渲染截图 → 多模态模型转写为结构化 Markdown。

    - 输出标记 content_format=markdown + headings，接入既有 Markdown 标题
      切分管线（标题路径 + 父子分块），保留标题/表格/图表语义；
    - 串行逐页解析，已解析标题层级注入下一页 prompt，跨页层级连贯；
    - 单页 VLM 失败降级为 PyMuPDF 文本提取，解析整体不中断；
    - 图片用区域原图（_extract_images_raw，防镜像/截断）挂页级供检索引用，
      不再把整页截图当图片入库。
    """
    client, vlm = _create_vlm_client(strategy)
    doc = fitz.open(file_path)
    pages = []
    raw_images = []
    heading_stack: List[dict] = []  # 跨页标题层级栈，供下一页 prompt 上下文
    try:
        # VLM 模式恒提取区域图（多模态解析的图片引用是核心能力，不受 extract_images 开关限制）
        all_images = _extract_images_raw(doc, doc_id)
        for page_idx in range(len(doc)):
            page = doc[page_idx]
            page_number = page_idx + 1
            images = [img for img in all_images if img["page_number"] == page_number]
            raw_images.extend(images)

            try:
                text = _vlm_transcribe_page(client, vlm, page, heading_stack)
                content_format = "markdown"
                headings = _extract_markdown_headings(text)
            except Exception as e:
                # 单页失败降级：该页回退 PyMuPDF 文本提取，不中断整体解析
                logger.warning("VLM 解析第 %s 页失败，降级为 PyMuPDF 文本提取: %s", page_number, e)
                text = page.get_text()
                content_format = "plain"
                headings = []

            # 维护跨页标题层级栈（与 split_pages_to_chunks 的栈逻辑一致）
            for h in headings:
                while heading_stack and heading_stack[-1]["level"] >= h["level"]:
                    heading_stack.pop()
                heading_stack.append({"level": h["level"], "text": h["text"]})

            pages.append(ParsedPage(
                page_number=page_number,
                text=text,
                images=[{"image_id": img["image_id"]} for img in images],
                tables=[],
                headings=headings,
                content_format=content_format,
            ))
    finally:
        doc.close()

    return ParseResult(
        doc_id=doc_id,
        pages=pages,
        raw_images=raw_images,
        mode_used="vlm",
    )


# ---------------------------------------------------------------------------
# PaddleOCR 解析（扫描件 / 位图文档）
# ---------------------------------------------------------------------------

_PADDLEOCR_ZOOM = 2.0  # PDF 页渲染位图的缩放倍数（~144 DPI）

# PaddleOCR 实例懒加载单例：模型大且初始化会触发模型下载，首次使用时才创建。
_paddle_ocr_instance = None


def _get_paddle_ocr():
    """返回懒加载的 PaddleOCR 单例。

    paddle/paddleocr 必须延迟导入：模块级导入会拖慢启动，且 Windows 上
    paddle 与 torch 存在 DLL 加载顺序冲突（先 paddle 后 torch 会导致 torch
    的 shm.dll 加载失败），因此这里固定先 import torch 再 import paddleocr。

    仅启用文本检测 + 识别（关闭文档方向分类/矫正/文本行方向分类），
    首次初始化时模型自动下载到 paddlex 默认缓存目录。
    """
    global _paddle_ocr_instance
    if _paddle_ocr_instance is None:
        import torch  # noqa: F401  # 必须先于 paddle 加载，规避 Windows DLL 冲突
        from paddleocr import PaddleOCR

        _paddle_ocr_instance = PaddleOCR(
            use_doc_orientation_classify=False,
            use_doc_unwarping=False,
            use_textline_orientation=False,
            lang="ch",
        )
    return _paddle_ocr_instance


def _pixmap_to_bgr_array(pix: "fitz.Pixmap"):
    """PyMuPDF Pixmap -> BGR np.ndarray（PaddleOCR ndarray 输入约定为 BGR 序）。"""
    import numpy as np

    arr = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)
    if pix.n == 1:  # 灰度 -> 三通道
        arr = np.repeat(arr, 3, axis=2)
    elif pix.n >= 4:  # 去 alpha
        arr = arr[..., :3]
    return arr[..., ::-1]  # RGB -> BGR


def _ocr_result_to_text(ocr_result) -> str:
    """将单页 OCR 结果按阅读顺序（从上到下、从左到右）拼接为纯文本。

    ocr_result 为 dict 结构（paddlex OCRResult 或等价的 mock）：
    rec_texts: List[str]，rec_boxes: [x1, y1, x2, y2] 列表（可选，缺失时保持原顺序）。
    """
    rec_texts = list(ocr_result.get("rec_texts") or [])
    if not rec_texts:
        return ""
    rec_boxes = ocr_result.get("rec_boxes")
    lines = []
    for idx, text in enumerate(rec_texts):
        x = y = 0.0
        try:
            if rec_boxes is not None and idx < len(rec_boxes):
                box = rec_boxes[idx]
                x, y = float(box[0]), float(box[1])
        except Exception:
            x = y = 0.0
        lines.append((y, x, idx, str(text).strip()))
    lines.sort(key=lambda item: (item[0], item[1], item[2]))
    return "\n".join(item[3] for item in lines if item[3])


def _ocr_one_image(image) -> str:
    """对单张图像（BGR ndarray 或图片文件路径）执行 OCR，返回拼接文本。"""
    ocr = _get_paddle_ocr()
    results = ocr.predict([image]) or []
    # 用量埋点：本地推理无 token，以图像张数（调用次数）计量
    usage_service.track("ocr", "paddleocr", scene="parse", extra={"image_count": 1})
    if not results:
        return ""
    return _ocr_result_to_text(results[0])


def _parse_with_paddleocr(file_path: str, doc_id: str, strategy: ParseStrategyConfig) -> ParseResult:
    """PaddleOCR 解析 PDF：每页渲染为位图（zoom 2x）后 OCR，按阅读顺序拼接文本。

    单页 OCR 失败时记日志并保留空 text，不中断整篇。
    输出结构与 pymupdf 路径一致；基础 OCR 不含版面/表格分析，tables/headings 为空。
    """
    doc = fitz.open(file_path)
    pages = []
    raw_images = []
    try:
        all_images = _extract_images_raw(doc, doc_id) if strategy.extract_images else []
        for page_idx in range(len(doc)):
            page = doc[page_idx]
            page_number = page_idx + 1

            page_text = ""
            try:
                mat = fitz.Matrix(_PADDLEOCR_ZOOM, _PADDLEOCR_ZOOM)
                pix = page.get_pixmap(matrix=mat)
                page_text = _ocr_one_image(_pixmap_to_bgr_array(pix))
            except Exception:
                logger.exception("PaddleOCR 解析失败: doc=%s page=%d", doc_id, page_number)

            images = [img for img in all_images if img["page_number"] == page_number]
            if images:
                placeholders = "\n".join(f"[IMG:{img['image_id']}]" for img in images)
                page_text = f"{page_text}\n\n{placeholders}" if page_text else placeholders
                raw_images.extend(images)

            pages.append(ParsedPage(
                page_number=page_number,
                text=page_text,
                images=[{"image_id": img["image_id"]} for img in images],
                tables=[],
                headings=[],
            ))
    finally:
        doc.close()

    return ParseResult(
        doc_id=doc_id,
        pages=pages,
        raw_images=raw_images,
        mode_used=strategy.parse_mode,
    )


def _parse_image_with_paddleocr(file_path: str, doc_id: str, strategy: ParseStrategyConfig) -> ParseResult:
    """单张图片直接 OCR（paddleocr/ocr 模式）。OCR 失败时记日志并返回空 text。"""
    try:
        text = _ocr_one_image(file_path)
    except Exception:
        logger.exception("PaddleOCR 解析图片失败: doc=%s file=%s", doc_id, file_path)
        text = ""

    with open(file_path, "rb") as f:
        image_data = f.read()
    ext = Path(file_path).suffix.lstrip(".") or "png"
    image_id = f"{doc_id}_img_01"
    raw_images = [{
        "image_id": image_id,
        "data": image_data,
        "content_type": f"image/{ext}",
        "width": 0,
        "height": 0,
        "page_number": 1,
        "seq": 1,
        "filename": f"{image_id}.{ext}",
    }]

    return ParseResult(
        doc_id=doc_id,
        pages=[ParsedPage(page_number=1, text=text, images=[{"image_id": image_id}])],
        raw_images=raw_images,
        mode_used=strategy.parse_mode,
    )


def _parse_image(file_path: str, doc_id: str, strategy: ParseStrategyConfig) -> ParseResult:
    """单张图片解析：paddleocr/ocr 模式直接 OCR，其余模式 fallback 到 VLM 描述。"""
    if strategy.parse_mode in ("paddleocr", "ocr"):
        return _parse_image_with_paddleocr(file_path, doc_id, strategy)
    # 当前实现：将图片作为一页，用 VLM 分析
    client, vlm = _create_vlm_client(strategy)
    with open(file_path, "rb") as f:
        image_data = f.read()

    import base64
    image_b64 = base64.b64encode(image_data).decode()
    ext = Path(file_path).suffix.lstrip(".") or "png"
    response = client.chat.completions.create(
        model=vlm["model"],
        messages=[{
            "role": "user",
            "content": [
                {"type": "text", "text": "请详细描述这张图片的内容，提取所有可见文字。"},
                {"type": "image_url", "image_url": {"url": f"data:image/{ext};base64,{image_b64}", "detail": vlm["detail"]}},
            ],
        }],
        max_tokens=vlm["max_tokens"],
    )
    _track_vlm_usage(response, vlm["model"])
    text = response.choices[0].message.content or ""

    image_id = f"{doc_id}_img_01"
    raw_images = [{
        "image_id": image_id,
        "data": image_data,
        "content_type": f"image/{ext}",
        "width": 0,
        "height": 0,
        "page_number": 1,
        "seq": 1,
        "filename": f"{image_id}.{ext}",
    }]

    return ParseResult(
        doc_id=doc_id,
        pages=[ParsedPage(page_number=1, text=text, images=[{"image_id": image_id}])],
        raw_images=raw_images,
        mode_used="vlm",
    )
